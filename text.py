import moviepy.editor as mp
from moviepy.video.io.ffmpeg_tools import ffmpeg_extract_subclip
import speech_recognition as sr 
from docx import Document

videoIndex = 1
nvideo = 0

while(videoIndex <= nvideo):
    # you can use these two instruction to shrink the video and work to a shorter clip. remeber to change the various paths
    # ffmpeg_extract_subclip("VIDEO/video (" + str(videoIndex) +").mp4", 0, 60, targetname="speech_video/new_video" + str(videoIndex) + ".mp4")
    # video = mp.VideoFileClip(r"temp_video/new_video" + str(videoIndex) + ".mp4")
    video = mp.VideoFileClip(r"VIDEO/video (" + str(videoIndex) + ").mp4") # you can remove this two instructions and change the various paths if you want to work only with audio
    video.audio.write_audiofile(r'temp_audio/converted_' + str(videoIndex) + '.wav')
    r = sr.Recognizer()
    audio = sr.AudioFile('temp_audio/converted_' + str(videoIndex) + '.wav')
    with audio as source:
        audio_file = r.record(source)
        result = r.recognize_google(audio_file, language='it-IT')
    if(str(result) != '[]'): # if the transcrition fails you will have a track of what failed
        documentPositive = Document()  
        documentPositive.add_paragraph(str(result))   
        documentPositive.add_page_break()    
        documentPositive.save('TEXT_SPEECH/recognized_' + str(videoIndex) + '.docx')
    else:
        documentNegative = Document()
        documentNegative.add_paragraph(str(result))   
        documentNegative.add_page_break()
        documentNegative.save('TEXT_SPEECH/NOT_RECOGNIZED_VIDNUMBER_' + str(videoIndex) + '.docx')
    videoIndex += 1
